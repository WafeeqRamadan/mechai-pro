# -*- coding: utf-8 -*-
"""
MechAI Pro v31-v36 — Universal Mechanical Intelligence Platform
========================================================================================
Knowledge-first Mechanical Engineering OS with:
- Supabase persistent storage and database
- Reference upload with metadata, extraction, chunking, duplicate detection, quality scoring
- Project/workspace-level retrieval with citations
- Supabase Auth login/signup
- Users, workspaces, projects, roles and basic permissions

Full-file build: Supabase Auth, persistent Reference Vault, project memory, Mechanical Knowledge Graph, multi-agent review board, calculators, CAD/SolidWorks and simulation intelligence. No OpenAI/Gemini dependency.
"""
from __future__ import annotations

import csv
import hashlib
import html
import io
import json
import math
import mimetypes
import os
import re
import textwrap
import uuid
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import streamlit as st

try:
    from supabase import Client, create_client  # type: ignore
except Exception:
    Client = None  # type: ignore
    create_client = None  # type: ignore

try:
    import PyPDF2  # type: ignore
except Exception:
    PyPDF2 = None

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
except Exception:
    A4 = None
    canvas = None

try:
    from openpyxl import Workbook  # type: ignore
except Exception:
    Workbook = None

APP_VERSION = "v36_1_ZERO_FRICTION_PROJECT_START_2026_06_15"
APP_TITLE = "MechAI Pro"
ROOT = Path(__file__).resolve().parent
GLOBAL_KNOWLEDGE_DIR = ROOT / "knowledge_packs"
LOCAL_CACHE_DIR = ROOT / "mechai_local_cache"
LOCAL_CACHE_DIR.mkdir(exist_ok=True)
GLOBAL_KNOWLEDGE_DIR.mkdir(exist_ok=True)

WORKSPACES = {
    "General engineering": {"icon": "🧠", "folder": None, "agent": "Chief Mechanical Engineering Board"},
    "Product R&D / Design": {"icon": "🛠️", "folder": "mechanical_design", "agent": "Mechanical Design Scientist"},
    "CAD / SolidWorks": {"icon": "🧩", "folder": "cad_solidworks", "agent": "CAD / SolidWorks Automation Scientist"},
    "Simulation / FEA": {"icon": "📊", "folder": "simulation_fea", "agent": "FEA Simulation Scientist"},
    "CFD / Thermal": {"icon": "🌊", "folder": "cfd_thermal", "agent": "CFD / Thermal Scientist"},
    "Manufacturing / DFM": {"icon": "🏭", "folder": "manufacturing_dfm", "agent": "Manufacturing DFM/DFA Scientist"},
    "Materials Selection": {"icon": "🧪", "folder": "materials_selection", "agent": "Materials Selection Scientist"},
    "Innovation / Patent": {"icon": "💡", "folder": "innovation_patent", "agent": "Innovation / Patent Scientist"},
}

PROJECT_TYPES = [
    "Product design", "DFM review", "Material selection", "FEA setup", "CFD setup",
    "SolidWorks automation", "Failure analysis", "Cost reduction", "Patent / innovation review",
]
SOURCE_TYPES = [
    "Open reference", "Personal reference", "Project reference", "Team reference", "Public datasheet",
    "Supplier catalog", "Design guide", "Standards summary", "Own engineering note", "Company standard",
]
CONFIDENTIALITY = ["Public", "Internal", "Private", "Confidential"]
APPROVAL_STATUS = ["Draft", "Under review", "Approved", "Deprecated"]
ROLES = ["Owner", "Admin", "Engineer", "Reviewer", "Viewer"]
UPLOAD_EXTENSIONS = {".pdf", ".txt", ".md", ".csv"}

# -----------------------------------------------------------------------------
# Styling
# -----------------------------------------------------------------------------
st.set_page_config(page_title="MechAI Pro", page_icon="⚙️", layout="wide", initial_sidebar_state="expanded")
st.markdown(
    """
<style>
:root { --bg: #0f0f0f; --panel: #171717; --panel2:#202020; --border:#2d2d2d; --text:#f1f1f1; --muted:#a3a3a3; }
html, body, [data-testid="stAppViewContainer"] { background:#0f0f0f; color:#f1f1f1; }
[data-testid="stSidebar"] { background:#151515; border-right:1px solid #2b2b2b; min-width:300px !important; max-width:300px !important; }
[data-testid="stHeader"] { background: rgba(15,15,15,0.85); }
.block-container { max-width: 1160px; padding-top: 1.2rem; padding-bottom: 7rem; }
.mech-card { background:#171717; border:1px solid #2d2d2d; border-radius:18px; padding:18px 20px; margin:10px 0; }
.mech-small { color:#a3a3a3; font-size:0.88rem; }
.mech-pill { display:inline-block; border:1px solid #343434; background:#1d1d1d; border-radius:999px; padding:4px 10px; margin:3px 4px 3px 0; font-size:0.82rem; color:#d0d0d0; }
.mech-risk-high { color:#ff8a8a; font-weight:700; }
.mech-risk-medium { color:#ffd27a; font-weight:700; }
.mech-risk-low { color:#8ee6a8; font-weight:700; }
.mech-source { border-left:3px solid #666; padding-left:12px; color:#cfcfcf; font-size:0.9rem; }
footer {visibility:hidden;}
#MainMenu {visibility:hidden;}
</style>
""",
    unsafe_allow_html=True,
)

# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------

def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def safe_slug(text: str, fallback: str = "item") -> str:
    text = (text or fallback).strip().lower()
    text = re.sub(r"[^a-z0-9\u0600-\u06FF]+", "_", text, flags=re.I).strip("_")
    return text[:80] or fallback


def get_secret(name: str, default: str = "") -> str:
    """Read a secret from Streamlit or environment and normalize whitespace/URL slashes."""
    value = default
    try:
        if name in st.secrets:
            value = str(st.secrets[name])
        else:
            value = os.environ.get(name, default)
    except Exception:
        value = os.environ.get(name, default)
    value = str(value).strip()
    if name == "SUPABASE_URL":
        value = value.rstrip("/")
    return value


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def estimate_tokens(text: str) -> int:
    return max(1, int(len(text.split()) * 1.25))


def read_uploaded_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        if PyPDF2 is None:
            return "[PDF extraction unavailable: PyPDF2 is not installed.]"
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            pages: List[str] = []
            for i, page in enumerate(reader.pages[:80]):
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    pages.append("")
            return "\n\n".join(pages).strip()
        except Exception as exc:
            return f"[PDF extraction failed: {exc}]"
    if suffix == ".csv":
        try:
            decoded = data.decode("utf-8", errors="replace")
            return decoded[:300_000]
        except Exception:
            return "[CSV extraction failed.]"
    try:
        return data.decode("utf-8", errors="replace")[:400_000]
    except Exception:
        return "[Text extraction failed.]"


def chunk_text(text: str, max_words: int = 420, overlap: int = 70) -> List[str]:
    words = re.findall(r"\S+", text or "")
    if not words:
        return []
    chunks: List[str] = []
    start = 0
    while start < len(words):
        end = min(start + max_words, len(words))
        chunk = " ".join(words[start:end]).strip()
        if len(chunk) > 80:
            chunks.append(chunk)
        if end >= len(words):
            break
        start = max(0, end - overlap)
    return chunks[:160]


def normalize_terms(text: str) -> List[str]:
    raw = re.findall(r"[A-Za-z0-9\u0600-\u06FF\+\-\.]+", (text or "").lower())
    stop = set("the a an and or of for to in on at by with from as is are be this that create make generate review setup plan".split())
    return [t for t in raw if len(t) > 2 and t not in stop]


QUERY_EXPANSIONS = {
    "plastic": ["thermoplastic", "injection", "molding", "abs", "pc", "pp", "wall", "draft", "rib", "boss", "warpage"],
    "cover": ["enclosure", "shell", "housing", "snap", "boss", "rib", "wall", "assembly"],
    "injection": ["molding", "draft", "gate", "ejector", "shrinkage", "sink", "warpage", "boss", "rib"],
    "dfm": ["manufacturing", "tooling", "process", "capability", "tolerance", "assembly", "cycle", "scrap"],
    "solidworks": ["macro", "vba", "step", "dxf", "drawing", "bom", "export", "assembly"],
    "fea": ["mesh", "boundary", "constraint", "load", "contacts", "element", "convergence", "validation"],
    "cfd": ["reynolds", "turbulence", "boundary", "mesh", "y+", "pressure", "heat", "convergence"],
    "material": ["density", "strength", "stiffness", "toughness", "temperature", "corrosion", "cost"],
}


def expanded_query_terms(query: str, project: Optional[Dict[str, Any]] = None) -> List[str]:
    base = normalize_terms(query)
    meta = ""
    if project:
        meta = " ".join(str(project.get(k) or "") for k in ["name", "project_type", "part_type", "material", "process", "manufacturing_method", "target_use"])
        base += normalize_terms(meta)
    expanded = list(base)
    for term in list(base):
        expanded += QUERY_EXPANSIONS.get(term, [])
    return list(dict.fromkeys(expanded))


def score_text(query_terms: List[str], text: str, source_quality: float = 0.6, workspace_boost: float = 1.0) -> float:
    lower = (text or "").lower()
    score = 0.0
    for term in query_terms:
        if term in lower:
            score += 1.0 + min(2.0, lower.count(term) * 0.1)
    return score * workspace_boost * (0.5 + source_quality)


def source_quality_score(source_type: str, approval_status: str, confidentiality: str, legal_note: str) -> float:
    score = 0.55
    if source_type in {"Standards summary", "Design guide", "Public datasheet", "Company standard", "Own engineering note"}:
        score += 0.15
    if approval_status == "Approved":
        score += 0.20
    elif approval_status == "Under review":
        score += 0.08
    elif approval_status == "Deprecated":
        score -= 0.25
    if legal_note and len(legal_note.strip()) > 15:
        score += 0.05
    if confidentiality in {"Private", "Confidential"}:
        score += 0.03  # not higher truth, but often project-specific
    return round(max(0.05, min(1.0, score)), 2)


def route_workspace(prompt: str, selected: str) -> str:
    p = prompt.lower()
    if any(x in p for x in ["solidworks", "macro", "vba", "step", "dxf", "bom", "drawing"]):
        return "CAD / SolidWorks"
    if any(x in p for x in ["fea", "ansys", "static structural", "modal", "buckling", "mesh", "element"]):
        return "Simulation / FEA"
    if any(x in p for x in ["cfd", "fluent", "reynolds", "turbulence", "y+", "pressure drop", "flow", "thermal"]):
        return "CFD / Thermal"
    if any(x in p for x in ["dfm", "dfa", "manufacturing", "injection", "molding", "machining", "sheet metal", "assembly", "tolerance"]):
        return "Manufacturing / DFM"
    if any(x in p for x in ["material", "abs", "pc", "pp", "steel", "aluminum", "elastomer", "corrosion"]):
        return "Materials Selection"
    if any(x in p for x in ["patent", "claim", "prior art", "invention", "novelty"]):
        return "Innovation / Patent"
    if selected != "General engineering":
        return selected
    return "Product R&D / Design"

# -----------------------------------------------------------------------------
# Supabase backend
# -----------------------------------------------------------------------------

@dataclass
class BackendStatus:
    configured: bool
    url_present: bool
    anon_present: bool
    service_present: bool
    bucket: str
    message: str


class SupabaseBackend:
    def __init__(self):
        self.url = get_secret("SUPABASE_URL")
        self.anon = get_secret("SUPABASE_ANON_KEY") or get_secret("SUPABASE_PUBLISHABLE_KEY")
        self.service = get_secret("SUPABASE_SERVICE_ROLE_KEY") or get_secret("SUPABASE_SECRET_KEY")
        self.bucket = get_secret("SUPABASE_BUCKET_REFERENCES", "mechai-references")
        self.auth_client: Optional[Client] = None
        self.db: Optional[Client] = None
        if create_client and self.url and self.anon:
            try:
                self.auth_client = create_client(self.url, self.anon)
                self.db = create_client(self.url, self.service or self.anon)
            except Exception as exc:
                st.sidebar.error(f"Supabase client error: {exc}")

    def status(self) -> BackendStatus:
        configured = bool(self.auth_client and self.db and self.url and self.anon)
        return BackendStatus(
            configured=configured,
            url_present=bool(self.url),
            anon_present=bool(self.anon),
            service_present=bool(self.service),
            bucket=self.bucket,
            message="Connected" if configured else "Not configured",
        )

    def sign_in(self, email: str, password: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        if not self.auth_client:
            return False, "Supabase Auth is not configured.", None
        try:
            res = self.auth_client.auth.sign_in_with_password({"email": email, "password": password})
            user = getattr(res, "user", None)
            session = getattr(res, "session", None)
            if not user:
                return False, "Sign-in failed.", None
            return True, "Signed in.", {"id": user.id, "email": user.email, "session": session}
        except Exception as exc:
            return False, f"Sign-in error: {exc}", None

    def sign_up(self, email: str, password: str) -> Tuple[bool, str]:
        if not self.auth_client:
            return False, "Supabase Auth is not configured."
        try:
            clean_email = email.strip()
            clean_password = password.strip()
            payload = {
                "email": clean_email,
                "password": clean_password,
                "options": {"email_redirect_to": "http://localhost:8501"},
            }
            res = self.auth_client.auth.sign_up(payload)
            user = getattr(res, "user", None)
            if user:
                return True, "Account created. Please sign in. If email confirmation is enabled, confirm your email first."
            return True, "Account request sent. Check Supabase email confirmation settings, then sign in."
        except Exception as exc:
            msg = str(exc)
            if "rate limit" in msg.lower():
                return False, "Supabase signup rate limit exceeded. Use Sign in if the account already exists, or wait a few minutes before creating another account."
            return False, f"Sign-up error: {msg}"

    def table(self, name: str):
        if not self.db:
            raise RuntimeError("Supabase database is not configured")
        return self.db.table(name)

    def safe_execute(self, query, default: Any = None) -> Any:
        try:
            res = query.execute()
            return getattr(res, "data", default)
        except Exception as exc:
            st.warning(f"Database operation failed: {exc}")
            return default

    # ---- Workspaces and members ----
    def create_workspace(self, name: str, workspace_type: str, user: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        payload = {
            "name": name.strip(),
            "workspace_type": workspace_type,
            "owner_user_id": user["id"],
            "created_at": now_iso(),
        }
        try:
            res = self.table("mechai_workspaces").insert(payload).execute()
            row = res.data[0]
            self.table("mechai_workspace_members").insert({
                "workspace_id": row["id"], "user_id": user["id"], "email": user["email"], "role": "Owner", "created_at": now_iso()
            }).execute()
            return row
        except Exception as exc:
            st.error(f"Create workspace failed: {exc}")
            return None

    def list_workspaces_for_user(self, user: Dict[str, Any]) -> List[Dict[str, Any]]:
        if not self.db or not user:
            return []
        rows_by_owner: List[Dict[str, Any]] = []
        rows_by_member: List[Dict[str, Any]] = []
        try:
            rows_by_owner = self.table("mechai_workspaces").select("*").eq("owner_user_id", user["id"]).execute().data or []
        except Exception:
            rows_by_owner = []
        try:
            mems = self.table("mechai_workspace_members").select("workspace_id,role,email,user_id").eq("email", user["email"]).execute().data or []
            ids = [m["workspace_id"] for m in mems]
            if ids:
                rows_by_member = self.table("mechai_workspaces").select("*").in_("id", ids).execute().data or []
        except Exception:
            rows_by_member = []
        merged = {r["id"]: r for r in rows_by_owner + rows_by_member}
        return list(merged.values())

    def get_role(self, workspace_id: str, user: Dict[str, Any]) -> str:
        try:
            if user and workspace_id:
                rows = self.table("mechai_workspace_members").select("role").eq("workspace_id", workspace_id).eq("email", user["email"]).limit(1).execute().data or []
                if rows:
                    return rows[0].get("role") or "Viewer"
        except Exception:
            pass
        return "Viewer"

    def add_member(self, workspace_id: str, email: str, role: str) -> Tuple[bool, str]:
        try:
            self.table("mechai_workspace_members").insert({
                "workspace_id": workspace_id, "email": email.strip().lower(), "role": role, "created_at": now_iso()
            }).execute()
            return True, "Member added. They will see the workspace after signing in with that email."
        except Exception as exc:
            return False, f"Add member failed: {exc}"

    def list_members(self, workspace_id: str) -> List[Dict[str, Any]]:
        try:
            return self.table("mechai_workspace_members").select("*").eq("workspace_id", workspace_id).execute().data or []
        except Exception:
            return []

    # ---- Projects ----
    def list_projects(self, workspace_id: str) -> List[Dict[str, Any]]:
        try:
            return self.table("mechai_projects").select("*").eq("workspace_id", workspace_id).order("created_at", desc=True).execute().data or []
        except Exception:
            return []

    def create_project(self, workspace_id: str, user: Dict[str, Any], payload: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        row = dict(payload)
        row.update({
            "workspace_id": workspace_id,
            "owner_user_id": user["id"],
            "created_at": now_iso(),
            "updated_at": now_iso(),
        })
        try:
            res = self.table("mechai_projects").insert(row).execute()
            return res.data[0]
        except Exception as exc:
            st.error(f"Create project failed: {exc}")
            return None

    def update_project(self, project_id: str, payload: Dict[str, Any]) -> bool:
        payload = dict(payload)
        payload["updated_at"] = now_iso()
        try:
            self.table("mechai_projects").update(payload).eq("id", project_id).execute()
            return True
        except Exception as exc:
            st.error(f"Update project failed: {exc}")
            return False

    # ---- References ----
    def find_duplicate(self, workspace_id: str, file_hash: str) -> Optional[Dict[str, Any]]:
        try:
            rows = self.table("mechai_references").select("*").eq("workspace_id", workspace_id).eq("file_hash", file_hash).limit(1).execute().data or []
            return rows[0] if rows else None
        except Exception:
            return None

    def upload_reference_file(self, storage_path: str, data: bytes, content_type: str) -> Tuple[bool, str]:
        if not self.db:
            return False, "Supabase is not configured."
        try:
            self.db.storage.from_(self.bucket).upload(
                storage_path,
                data,
                file_options={"content-type": content_type, "x-upsert": "true"},
            )
            return True, storage_path
        except Exception as exc:
            # Try update if upload complains that file exists.
            try:
                self.db.storage.from_(self.bucket).update(storage_path, data, file_options={"content-type": content_type})
                return True, storage_path
            except Exception:
                return False, f"Storage upload failed: {exc}"

    def insert_reference(self, ref: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        try:
            res = self.table("mechai_references").insert(ref).execute()
            return res.data[0]
        except Exception as exc:
            st.error(f"Reference metadata insert failed: {exc}")
            return None

    def insert_chunks(self, chunks: List[Dict[str, Any]]) -> bool:
        if not chunks:
            return True
        try:
            self.table("mechai_reference_chunks").insert(chunks).execute()
            return True
        except Exception as exc:
            st.error(f"Chunk insert failed: {exc}")
            return False

    def list_references(self, workspace_id: str, project_id: Optional[str] = None, search: str = "") -> List[Dict[str, Any]]:
        try:
            q = self.table("mechai_references").select("*").eq("workspace_id", workspace_id).order("created_at", desc=True)
            if project_id:
                # include project-specific and workspace/global refs
                rows = q.execute().data or []
                rows = [r for r in rows if not r.get("project_id") or r.get("project_id") == project_id]
            else:
                rows = q.execute().data or []
            if search.strip():
                terms = normalize_terms(search)
                rows = [r for r in rows if any(t in json.dumps(r, ensure_ascii=False).lower() for t in terms)]
            return rows
        except Exception:
            return []

    def search_reference_chunks(self, workspace_id: str, project_id: Optional[str], query: str, project: Optional[Dict[str, Any]], top_k: int = 7) -> List[Dict[str, Any]]:
        try:
            rows = self.table("mechai_reference_chunks").select("*, mechai_references(title,source_type,confidentiality,approval_status,revision,tags,file_name,source_quality)").eq("workspace_id", workspace_id).limit(2000).execute().data or []
        except Exception:
            return []
        if project_id:
            rows = [r for r in rows if not r.get("project_id") or r.get("project_id") == project_id]
        terms = expanded_query_terms(query, project)
        scored: List[Dict[str, Any]] = []
        for r in rows:
            ref_meta = r.get("mechai_references") or {}
            quality = float(ref_meta.get("source_quality") or r.get("source_quality") or 0.6)
            score = score_text(terms, r.get("chunk_text") or "", quality, 1.1)
            # Workspace and project-specific boosts
            if r.get("project_id") == project_id:
                score *= 1.35
            if score > 0:
                item = dict(r)
                item["score"] = round(score, 3)
                item["source_title"] = ref_meta.get("title") or "Uploaded reference"
                item["source_file"] = ref_meta.get("file_name") or "reference"
                item["source_type"] = ref_meta.get("source_type") or "Reference"
                item["approval_status"] = ref_meta.get("approval_status") or "Draft"
                item["source_quality"] = quality
                scored.append(item)
        scored.sort(key=lambda x: x.get("score", 0), reverse=True)
        return scored[:top_k]

    # ---- Memory ----
    def add_memory(self, workspace_id: str, project_id: str, user: Dict[str, Any], memory_type: str, title: str, content: str, metadata: Optional[Dict[str, Any]] = None) -> None:
        try:
            self.table("mechai_project_memory").insert({
                "workspace_id": workspace_id,
                "project_id": project_id,
                "user_id": user.get("id") if user else None,
                "memory_type": memory_type,
                "title": title[:200],
                "content": content[:10000],
                "metadata": metadata or {},
                "created_at": now_iso(),
            }).execute()
        except Exception:
            pass

    def list_memory(self, workspace_id: str, project_id: str, limit: int = 100) -> List[Dict[str, Any]]:
        try:
            return self.table("mechai_project_memory").select("*").eq("workspace_id", workspace_id).eq("project_id", project_id).order("created_at", desc=True).limit(limit).execute().data or []
        except Exception:
            return []


backend = SupabaseBackend()

# -----------------------------------------------------------------------------
# Global knowledge retrieval
# -----------------------------------------------------------------------------

def seed_minimum_global_knowledge() -> None:
    seed = {
        "manufacturing_dfm/injection_molding_expert.md": """# Injection Molding Expert Pack\nScope: plastic part DFM, wall thickness, draft, ribs, bosses, gate, ejection, sink, warpage, tolerance and quality.\nRequired inputs: material grade, nominal wall, part size, ribs, bosses, surface class, annual volume, tolerance targets.\nRules: use uniform wall thickness; add draft; avoid thick bosses; ribs should support stiffness without sink; gate/parting/ejection must be considered before tooling; validate with mold-flow or first article for critical parts.\n""",
        "manufacturing_dfm/tolerance_capability_expert.md": """# Tolerance Capability Pack\nScope: CTQ dimensions, process capability, inspection and supplier feasibility.\nRules: tight tolerances must be justified by function; process capability must be matched to tolerance; unknown CTQ lowers confidence.\n""",
        "materials_selection/thermoplastics.md": """# Thermoplastics Pack\nScope: ABS, PC, PP, PA, POM and plastic material selection.\nRules: choose by stiffness, toughness, temperature, chemical exposure, appearance, process compatibility, cost, shrinkage and availability.\n""",
        "simulation_fea/static_structural.md": """# Static Structural FEA Pack\nScope: loads, constraints, contacts, material, mesh convergence, validation and failure criteria.\nRules: define objective, load path and constraints before meshing; validate with hand calculations or test data.\n""",
        "cfd_thermal/reynolds_number.md": """# Reynolds Number Pack\nEquation: Re = rho*V*D/mu. Laminar pipe flow usually below about 2300, transitional around 2300-4000, turbulent above about 4000. Validate assumptions and properties.\n""",
        "cad_solidworks/macro_generation.md": """# SolidWorks Macro Generation Pack\nRules: create safe VBA macros with error handling, clear object model steps, file path warnings and no destructive actions without confirmation.\n""",
    }
    for rel, content in seed.items():
        p = GLOBAL_KNOWLEDGE_DIR / rel
        p.parent.mkdir(parents=True, exist_ok=True)
        if not p.exists():
            p.write_text(content, encoding="utf-8")


def load_global_knowledge_docs() -> List[Dict[str, Any]]:
    seed_minimum_global_knowledge()
    docs: List[Dict[str, Any]] = []
    for path in sorted(GLOBAL_KNOWLEDGE_DIR.rglob("*.md")):
        try:
            rel = path.relative_to(GLOBAL_KNOWLEDGE_DIR).as_posix()
            workspace_folder = rel.split("/")[0] if "/" in rel else "general"
            text = path.read_text(encoding="utf-8", errors="replace")
            for idx, chunk in enumerate(chunk_text(text, max_words=360, overlap=50)):
                docs.append({
                    "kind": "global",
                    "workspace_folder": workspace_folder,
                    "source_path": rel,
                    "chunk_index": idx,
                    "text": chunk,
                    "source_quality": 0.72,
                    "title": Path(rel).name,
                })
        except Exception:
            continue
    return docs


def search_global_knowledge(query: str, workspace: str, project: Optional[Dict[str, Any]], top_k: int = 8) -> List[Dict[str, Any]]:
    terms = expanded_query_terms(query, project)
    folder = WORKSPACES.get(workspace, {}).get("folder")
    docs = load_global_knowledge_docs()
    scored: List[Dict[str, Any]] = []
    for d in docs:
        boost = 1.0
        if folder and d.get("workspace_folder") == folder:
            boost = 1.45
        elif workspace == "General engineering":
            boost = 1.0
        score = score_text(terms, d["text"], float(d.get("source_quality", 0.7)), boost)
        if score > 0:
            item = dict(d)
            item["score"] = round(score, 3)
            scored.append(item)
    scored.sort(key=lambda x: x.get("score", 0), reverse=True)
    return scored[:top_k]



# -----------------------------------------------------------------------------
# v31-v36 Mechanical Intelligence Layer
# -----------------------------------------------------------------------------

GRAPH_LIBRARY = {
    "cover/enclosure": {
        "category": "part",
        "aliases": ["cover", "enclosure", "lid", "housing", "case", "shell"],
        "links": ["injection molding", "sheet metal", "ABS", "PC", "PP", "wall thickness", "draft", "ribs", "bosses", "sink marks", "warpage", "snap-fit", "cosmetic surface", "first article inspection"],
        "risks": ["sink marks", "warpage", "boss cracking", "snap-fit tolerance", "cosmetic defects", "assembly interference"],
        "validation": ["CAD section review", "tooling review", "first article inspection", "process capability", "fit/assembly test"],
    },
    "bracket": {
        "category": "part",
        "aliases": ["bracket", "mount", "support", "hanger"],
        "links": ["static structural", "bending", "stress concentration", "fillet", "bolt preload", "fatigue", "mesh convergence", "yield", "deflection"],
        "risks": ["yielding", "fatigue cracking", "bolt hole stress", "buckling", "excessive deflection"],
        "validation": ["hand calculation", "FEA mesh convergence", "load test", "torque/preload verification"],
    },
    "shaft": {
        "category": "part",
        "aliases": ["shaft", "spindle", "axle"],
        "links": ["torsion", "bending", "fatigue", "keyway", "bearing seat", "critical speed", "surface finish", "heat treatment"],
        "risks": ["torsional failure", "fatigue", "deflection", "bearing misalignment", "keyway stress concentration"],
        "validation": ["torsional stress", "fatigue factor of safety", "deflection check", "bearing life", "runout inspection"],
    },
    "pipe/duct": {
        "category": "part",
        "aliases": ["pipe", "duct", "channel", "manifold", "hose"],
        "links": ["Reynolds number", "pressure drop", "internal flow", "turbulence", "heat transfer", "roughness", "minor losses"],
        "risks": ["pressure drop too high", "noise", "erosion", "thermal loss", "flow separation"],
        "validation": ["Reynolds calculation", "pressure-drop estimate", "flow test", "mass balance"],
    },
    "injection molding": {
        "category": "process",
        "aliases": ["injection", "mold", "mould", "plastic molding", "injection molded"],
        "links": ["wall thickness", "draft", "ribs", "bosses", "gate", "ejector", "parting line", "shrinkage", "sink marks", "warpage", "cycle time", "tooling"],
        "risks": ["sink marks", "warpage", "short shot", "flash", "weld lines", "ejection marks", "tooling complexity"],
        "validation": ["DFM review", "mold-flow review", "tooling review", "FAI", "Cp/Cpk"],
    },
    "sheet metal": {
        "category": "process",
        "aliases": ["sheet metal", "bend", "laser cut", "press brake", "flat pattern"],
        "links": ["bend allowance", "k-factor", "minimum flange", "relief", "grain direction", "springback", "DXF"],
        "risks": ["cracking", "springback", "wrong flat pattern", "tool collision", "burrs"],
        "validation": ["bend allowance check", "flat pattern review", "first-off bend inspection"],
    },
    "machining": {
        "category": "process",
        "aliases": ["machining", "milling", "turning", "cnc", "drilling"],
        "links": ["tool access", "tolerance", "surface finish", "workholding", "cycle time", "tool wear"],
        "risks": ["tool chatter", "difficult fixturing", "over-tight tolerance", "burrs", "distortion"],
        "validation": ["process plan", "fixture review", "first article inspection", "capability study"],
    },
    "ABS": {
        "category": "material",
        "aliases": ["abs", "acrylonitrile butadiene styrene"],
        "links": ["injection molding", "impact resistance", "cosmetic surface", "creep", "heat deflection", "UV aging"],
        "risks": ["thermal softening", "UV degradation", "chemical attack", "creep"],
        "validation": ["datasheet review", "temperature test", "impact test", "aging/UV check"],
    },
    "aluminum": {
        "category": "material",
        "aliases": ["aluminum", "aluminium", "6061", "6082", "7075"],
        "links": ["machining", "extrusion", "corrosion", "anodizing", "fatigue", "yield"],
        "risks": ["galvanic corrosion", "fatigue", "thread stripping", "distortion"],
        "validation": ["material certificate", "surface treatment review", "fatigue check"],
    },
    "steel": {
        "category": "material",
        "aliases": ["steel", "mild steel", "stainless", "carbon steel"],
        "links": ["welding", "machining", "fatigue", "corrosion", "heat treatment"],
        "risks": ["corrosion", "weld distortion", "fatigue", "brittle fracture"],
        "validation": ["material certificate", "coating review", "weld inspection", "fatigue check"],
    },
}

AGENT_BOARD = {
    "Design Engineer": "function, load path, geometry, interfaces, tolerances and failure modes",
    "Manufacturing Engineer": "process capability, tooling, cycle time, scrap and repeatable production",
    "Materials Engineer": "material grade, environment, strength, stiffness, creep, corrosion and availability",
    "Quality Engineer": "CTQs, inspection strategy, capability, FAI, acceptance criteria and release evidence",
    "Simulation Engineer": "physics, boundary conditions, validation, mesh/convergence and interpretation limits",
    "Cost Engineer": "cost drivers, part count, material utilization, cycle time, tooling and supplier risk",
    "Patent / Innovation Reviewer": "novelty framing, claimable features, prior-art risk and prototype evidence",
}

RELEASE_CRITERIA = {
    "Pass": "Inputs and evidence are mature enough for the next formal engineering stage.",
    "Conditional Pass": "Proceed only after closing high-priority missing inputs and verification actions.",
    "Engineering Hold": "Do not release; key engineering assumptions or evidence are missing.",
    "Stop / Redesign": "Fundamental design/process risk appears incompatible with the stated objective.",
}


def graph_match_entities(prompt: str, project: Optional[Dict[str, Any]] = None) -> List[str]:
    blob = " ".join([prompt or "", json.dumps(project or {}, ensure_ascii=False)]).lower()
    found: List[str] = []
    for node, data in GRAPH_LIBRARY.items():
        aliases = [node] + list(data.get("aliases", []))
        if any(str(a).lower() in blob for a in aliases):
            found.append(node)
    return found


def graph_context(prompt: str, project: Optional[Dict[str, Any]], workspace: str) -> Dict[str, Any]:
    entities = graph_match_entities(prompt, project)
    if not entities:
        # workspace-biased defaults keep the graph useful even with vague questions.
        if workspace == "Manufacturing / DFM":
            entities = ["injection molding"] if "mold" in prompt.lower() or "plastic" in prompt.lower() else ["machining"]
        elif workspace == "Simulation / FEA":
            entities = ["bracket"]
        elif workspace == "CFD / Thermal":
            entities = ["pipe/duct"]
        elif workspace == "CAD / SolidWorks":
            entities = ["cover/enclosure"]
    links: List[str] = []
    risks: List[str] = []
    validation: List[str] = []
    categories: Dict[str, List[str]] = {}
    for e in entities:
        d = GRAPH_LIBRARY.get(e, {})
        categories.setdefault(d.get("category", "concept"), []).append(e)
        links.extend(d.get("links", []))
        risks.extend(d.get("risks", []))
        validation.extend(d.get("validation", []))
    def unique(seq: Iterable[str]) -> List[str]:
        return list(dict.fromkeys([str(x) for x in seq if str(x).strip()]))
    return {
        "entities": unique(entities),
        "categories": categories,
        "linked_concepts": unique(links)[:16],
        "failure_modes": unique(risks)[:12],
        "validation_methods": unique(validation)[:12],
    }


def source_family_bias_from_graph(gctx: Dict[str, Any]) -> List[str]:
    concepts = " ".join(gctx.get("linked_concepts", []) + gctx.get("failure_modes", [])).lower()
    wanted: List[str] = []
    if any(x in concepts for x in ["wall", "draft", "shrink", "sink", "warpage", "tooling"]):
        wanted.extend(["injection_molding", "manufacturing_dfm", "thermoplastics", "tolerance"])
    if any(x in concepts for x in ["mesh", "boundary", "convergence", "stress", "yield"]):
        wanted.extend(["static_structural", "mesh_convergence", "validation", "simulation_fea"])
    if any(x in concepts for x in ["reynolds", "pressure", "turbulence", "heat"]):
        wanted.extend(["internal_flow", "reynolds", "pressure_drop", "heat_transfer", "cfd_thermal"])
    if any(x in concepts for x in ["macro", "drawing", "bom", "dxf", "step"]):
        wanted.extend(["macro", "solidworks", "cad_solidworks", "export"])
    return list(dict.fromkeys(wanted))


def board_review(prompt: str, workspace: str, project: Optional[Dict[str, Any]], missing: List[str], risks: List[Tuple[str, str, str]], gctx: Dict[str, Any]) -> List[Dict[str, str]]:
    missing_blob = ", ".join(missing) if missing else "none critical detected"
    high_risks = [r[0] for r in risks if r[1] == "High"]
    board: List[Dict[str, str]] = []
    for agent, focus in AGENT_BOARD.items():
        if workspace == "Manufacturing / DFM" and agent == "Simulation Engineer":
            stance = "Advisory"
        elif high_risks and agent in {"Quality Engineer", "Manufacturing Engineer"}:
            stance = "Concern"
        elif missing and agent in {"Design Engineer", "Materials Engineer"}:
            stance = "Needs input"
        else:
            stance = "Preliminary pass"
        if agent == "Design Engineer":
            note = f"Confirm function, interfaces and geometry before release. Missing input focus: {missing_blob}."
        elif agent == "Manufacturing Engineer":
            note = f"Process capability and tooling constraints dominate. Key graph risks: {', '.join(gctx.get('failure_modes', [])[:5]) or 'not enough evidence'}."
        elif agent == "Materials Engineer":
            note = "Material grade and environment must be explicit before final selection, simulation or process release."
        elif agent == "Quality Engineer":
            note = "Define CTQ dimensions, inspection method, acceptance criteria and evidence records before any release gate."
        elif agent == "Simulation Engineer":
            note = "Use simulation as evidence only with validated loads, constraints, material properties and convergence checks."
        elif agent == "Cost Engineer":
            note = "Rank changes by risk reduction, tooling impact, cycle time, material use and implementation effort."
        else:
            note = "If the concept is novel, separate technical novelty from manufacturability and commercial value."
        board.append({"agent": agent, "focus": focus, "stance": stance, "note": note})
    return board


def release_gate_decision(score: int, missing: List[str], risks: List[Tuple[str, str, str]], evidence_level: str) -> str:
    high = sum(1 for _, level, _ in risks if level == "High")
    unknown = sum(1 for _, level, _ in risks if level == "Unknown")
    if score < 45 or high >= 3:
        return "Engineering Hold"
    if score < 70 or missing or unknown:
        return "Conditional Pass"
    if "Level 4" in evidence_level and score >= 80:
        return "Pass"
    return "Conditional Pass"


def evidence_score(ref_sources: List[Dict[str, Any]], global_sources: List[Dict[str, Any]], calc_notes: List[str]) -> Tuple[str, int]:
    pts = 0
    if global_sources:
        pts += 20
    if ref_sources:
        pts += 35
    if calc_notes:
        pts += 25
    if len(global_sources) + len(ref_sources) >= 5:
        pts += 10
    pts = min(100, pts)
    if pts >= 75:
        level = "Level 4 — internally sourced + reference-backed + deterministic checks"
    elif pts >= 50:
        level = "Level 3 — internal knowledge with retrieved sources"
    elif pts >= 20:
        level = "Level 2 — internal qualitative engineering guidance"
    else:
        level = "Level 1 — preliminary frame only"
    return level, pts


def deterministic_checks_v36(text: str, project: Optional[Dict[str, Any]] = None) -> List[str]:
    notes: List[str] = []
    rn = calculate_reynolds_if_possible(text)
    if rn:
        notes.append(rn)
    p = text.lower()
    # Beam center point load rough parser: 2 kN, 500 mm, E/I if provided.
    if "beam" in p and any(x in p for x in ["load", "span", "deflection"]):
        notes.append("Beam validator available: provide P, span L, modulus E and second moment I to calculate δ = PL³/(48EI) and bending stress checks.")
    if "shaft" in p or "torque" in p:
        notes.append("Shaft validator available: provide torque, shaft diameter, material allowable shear, keyway details and bearing span to check torsion, fatigue and deflection.")
    if "bearing" in p:
        notes.append("Bearing validator available: provide dynamic rating C, equivalent load P, bearing type and speed to estimate L10 life.")
    if "wall" in p or "injection" in p or "mold" in p:
        notes.append("Injection molding validator: compare rib/boss thickness to nominal wall, check draft, gate/ejector feasibility, shrinkage and sink/warpage risks.")
    if "sheet" in p or "bend" in p:
        notes.append("Sheet-metal validator: provide bend angle, inside radius, thickness and K-factor to calculate bend allowance and flat pattern risk.")
    return list(dict.fromkeys(notes))[:8]


def knowledge_graph_markdown(gctx: Dict[str, Any]) -> str:
    lines = ["#### Mechanical Knowledge Graph context"]
    if gctx.get("entities"):
        lines.append("**Detected nodes:** " + ", ".join(gctx["entities"]))
    if gctx.get("linked_concepts"):
        lines.append("**Connected engineering concepts:** " + ", ".join(gctx["linked_concepts"][:12]))
    if gctx.get("failure_modes"):
        lines.append("**Failure/risk modes:** " + ", ".join(gctx["failure_modes"][:10]))
    if gctx.get("validation_methods"):
        lines.append("**Validation methods:** " + ", ".join(gctx["validation_methods"][:10]))
    return "\n\n".join(lines)


# -----------------------------------------------------------------------------
# Engineering reasoning and calculators
# -----------------------------------------------------------------------------

def detect_missing_inputs(prompt: str, project: Optional[Dict[str, Any]], workspace: str) -> List[str]:
    p = prompt.lower()
    missing: List[str] = []
    project = project or {}
    def has_field(k: str) -> bool:
        return bool(str(project.get(k) or "").strip())
    if workspace == "Manufacturing / DFM":
        for label, key in [("material grade", "material"), ("manufacturing process", "process"), ("annual volume", "annual_volume")]:
            if not has_field(key):
                missing.append(label)
        for item in ["nominal wall thickness", "CAD/STEP or drawing", "CTQ dimensions", "surface finish class"]:
            if item.lower() not in p:
                missing.append(item)
    elif workspace == "Simulation / FEA":
        for item in ["material properties", "load magnitude and direction", "constraints", "contacts", "mesh convergence criterion", "failure criterion"]:
            if item.split()[0].lower() not in p:
                missing.append(item)
    elif workspace == "CFD / Thermal":
        for item in ["fluid properties", "domain geometry", "inlet/outlet boundary conditions", "temperature/heat load", "mesh/y+ target", "validation method"]:
            if item.split()[0].lower() not in p:
                missing.append(item)
    elif workspace == "Materials Selection":
        for item in ["functional requirements", "temperature range", "chemical exposure", "manufacturing process", "cost target", "availability"]:
            if item.split()[0].lower() not in p:
                missing.append(item)
    elif workspace == "CAD / SolidWorks":
        for item in ["SolidWorks version", "file path strategy", "target document type", "export format", "overwrite policy"]:
            if item.split()[0].lower() not in p:
                missing.append(item)
    return list(dict.fromkeys(missing))[:10]


def maturity_score(missing: List[str]) -> Tuple[int, str]:
    score = max(15, 100 - len(missing) * 10)
    if score >= 80:
        label = "High — enough input for a stronger engineering recommendation"
    elif score >= 55:
        label = "Medium — useful preliminary review, but not release-ready"
    else:
        label = "Low — preliminary only; major assumptions remain"
    return score, label


def risk_matrix(prompt: str, workspace: str, missing: List[str]) -> List[Tuple[str, str, str]]:
    p = prompt.lower()
    rows: List[Tuple[str, str, str]] = []
    if workspace == "Manufacturing / DFM":
        rows = [
            ("Wall thickness", "High" if any("wall" in m for m in missing) else "Medium", "Unknown or non-uniform wall thickness drives sink, warpage and cycle-time risk."),
            ("Draft / ejection", "Medium", "Injection molded parts need release strategy; absent draft data blocks tooling confidence."),
            ("Ribs / bosses", "High" if "cover" in p or "enclosure" in p else "Medium", "Covers often need bosses, ribs and snap features that can create sink and stress concentrations."),
            ("Tooling strategy", "Medium", "Parting line, gate and ejector strategy are not defined."),
            ("Tolerance capability", "Unknown" if any("ctq" in m.lower() for m in missing) else "Medium", "CTQ dimensions and process capability are not yet established."),
            ("Assembly", "Unknown", "Mating parts and assembly method are not defined."),
        ]
    elif workspace == "Simulation / FEA":
        rows = [
            ("Boundary conditions", "High" if "constraints" in " ".join(missing).lower() else "Medium", "FEA result validity depends strongly on realistic constraints."),
            ("Loads", "High" if "load" in " ".join(missing).lower() else "Medium", "Load magnitude, direction and application area must be explicit."),
            ("Material model", "High" if "material" in " ".join(missing).lower() else "Medium", "Material properties and failure criterion affect stress interpretation."),
            ("Mesh convergence", "High", "A single mesh result is not enough for release decisions."),
            ("Validation", "High", "Hand calculation or test evidence is required for confidence."),
        ]
    elif workspace == "CFD / Thermal":
        rows = [
            ("Flow regime", "High", "Reynolds number and fluid properties are needed before model choice."),
            ("Boundary conditions", "High", "Inlet/outlet/thermal conditions determine solution meaning."),
            ("Mesh/y+", "Medium", "Turbulence and wall treatment require mesh strategy."),
            ("Convergence", "Medium", "Residuals must be checked with mass/energy balance."),
            ("Validation", "High", "Analytical or test comparison is needed."),
        ]
    else:
        rows = [("Input completeness", "Medium" if missing else "Low", "Review depends on the missing input list and source evidence level.")]
    return rows


def risk_level_to_points(level: str) -> int:
    return {"High": 22, "Medium": 12, "Low": 4, "Unknown": 16}.get(level, 10)


def score_from_risks(rows: List[Tuple[str, str, str]], missing: List[str]) -> int:
    penalty = sum(risk_level_to_points(r[1]) for r in rows) + len(missing) * 3
    return max(5, min(95, 100 - penalty))


def calculate_reynolds_if_possible(text: str) -> Optional[str]:
    p = text.lower()
    if "reynolds" not in p and not ("pipe" in p and any(x in p for x in ["water", "air", "flow"])):
        return None
    # Try to detect diameter in mm/m and velocity in m/s.
    d = None
    v = None
    m = re.search(r"(\d+(?:\.\d+)?)\s*mm", p)
    if m:
        d = float(m.group(1)) / 1000.0
    m = re.search(r"(\d+(?:\.\d+)?)\s*m/s", p)
    if m:
        v = float(m.group(1))
    rho, mu = 1000.0, 0.001
    if "air" in p:
        rho, mu = 1.225, 1.8e-5
    if d and v:
        re_val = rho * v * d / mu
        if re_val < 2300:
            regime = "laminar"
        elif re_val < 4000:
            regime = "transitional"
        else:
            regime = "turbulent"
        return f"Deterministic check — Reynolds number: Re = ρVD/μ = {re_val:,.0f}. Estimated regime: **{regime}**. Assumed {'air' if 'air' in p else 'water'} properties near room temperature."
    return "Reynolds check requested, but diameter and velocity were not both detected. Provide fluid, diameter, velocity and temperature."


def generate_vba_macro(request: str) -> str:
    return textwrap.dedent(f'''
        ' MechAI Pro generated SolidWorks VBA skeleton
        ' Request: {request.replace(chr(10), ' ')[:180]}
        ' Safety: run on copied files first. Review output paths before execution.
        Option Explicit

        Sub main()
            On Error GoTo EH
            Dim swApp As SldWorks.SldWorks
            Dim swModel As SldWorks.ModelDoc2
            Set swApp = Application.SldWorks
            Set swModel = swApp.ActiveDoc
            If swModel Is Nothing Then
                MsgBox "No active SolidWorks document.", vbExclamation
                Exit Sub
            End If

            Dim modelPath As String
            modelPath = swModel.GetPathName
            If modelPath = "" Then
                MsgBox "Save the document before export.", vbExclamation
                Exit Sub
            End If

            Dim outFolder As String
            outFolder = Left(modelPath, InStrRev(modelPath, "\\")) & "MechAI_Exports\\"
            If Dir(outFolder, vbDirectory) = "" Then MkDir outFolder

            Dim baseName As String
            baseName = Mid(modelPath, InStrRev(modelPath, "\\") + 1)
            baseName = Left(baseName, InStrRev(baseName, ".") - 1)

            Dim errs As Long, warns As Long
            swModel.Extension.SaveAs outFolder & baseName & ".STEP", 0, 0, Nothing, errs, warns
            swModel.Extension.SaveAs outFolder & baseName & ".DXF", 0, 0, Nothing, errs, warns

            MsgBox "Export completed: " & outFolder, vbInformation
            Exit Sub
        EH:
            MsgBox "Macro failed: " & Err.Description, vbCritical
        End Sub
    ''').strip()


def generate_apdl(request: str) -> str:
    return textwrap.dedent(f'''
        ! MechAI Pro ANSYS APDL starter
        ! Request: {request.replace(chr(10), ' ')[:180]}
        /PREP7
        ! TODO: define material properties
        ! MP,EX,1,2.1E11
        ! MP,PRXY,1,0.3
        ! TODO: import/create geometry and mesh
        ! TODO: apply boundary conditions and loads
        ! Use mesh convergence and hand-calculation validation before release.
        /SOLU
        ANTYPE,0
        ! SOLVE
        /POST1
        ! PRNSOL,S,COMP
    ''').strip()


def generate_fluent_journal(request: str) -> str:
    return textwrap.dedent(f'''
        ; MechAI Pro Fluent journal starter
        ; Request: {request.replace(chr(10), ' ')[:180]}
        ; TODO: read mesh, set units, define material, boundary conditions, turbulence model.
        /file/read-case "case.cas.h5"
        /solve/initialize/hyb-initialization
        ; TODO: set monitors for residuals, mass balance, pressure drop, temperature.
        /solve/iterate 500
        /file/write-case-data "mechai_result.cas.h5"
    ''').strip()


def build_engineering_answer(
    prompt: str,
    selected_workspace: str,
    project: Optional[Dict[str, Any]],
    global_sources: List[Dict[str, Any]],
    ref_sources: List[Dict[str, Any]],
) -> str:
    routed = route_workspace(prompt, selected_workspace)
    agent = WORKSPACES.get(routed, {}).get("agent", "Mechanical Engineering Board")
    gctx = graph_context(prompt, project, routed)
    # Re-rank global sources using graph-derived source family bias.
    graph_bias = source_family_bias_from_graph(gctx)
    if graph_bias and global_sources:
        def boosted_score(src: Dict[str, Any]) -> float:
            blob = f"{src.get('title','')} {src.get('source_path','')} {src.get('text','')[:500]}".lower()
            boost = sum(1.2 for b in graph_bias if b.lower() in blob)
            return float(src.get("score", 0)) + boost
        global_sources = sorted(global_sources, key=boosted_score, reverse=True)

    missing = detect_missing_inputs(prompt, project, routed)
    maturity, maturity_label = maturity_score(missing)
    risks = risk_matrix(prompt, routed, missing)
    base_score = score_from_risks(risks, missing)
    calc_notes = deterministic_checks_v36(prompt, project)
    evidence_level, evidence_pts = evidence_score(ref_sources, global_sources, calc_notes)
    gate = release_gate_decision(base_score, missing, risks, evidence_level)
    board = board_review(prompt, routed, project, missing, risks, gctx)

    project_lines = []
    if project:
        for label, key in [("Project", "name"), ("Type", "project_type"), ("Part type", "part_type"), ("Material", "material"), ("Process", "process"), ("Annual volume", "annual_volume"), ("Target use", "target_use")]:
            val = project.get(key)
            if val:
                project_lines.append(f"- {label}: {val}")

    out: List[str] = []
    out.append(f"### MechAI v31–v36 Mechanical Intelligence Platform — {routed}\n")
    out.append(f"**Primary board:** {agent}")
    out.append(f"**Mode:** Internal Knowledge + Persistent Reference Vault")
    out.append(f"**Knowledge graph:** v31 Mechanical Knowledge Graph active")
    out.append(f"**Agent board:** v32 Multi-agent mechanical review active")
    out.append(f"**Review board:** v33 Release gate + evidence evaluation active")
    out.append(f"**Calculators:** v34 deterministic validators available")
    out.append(f"**CAD/Simulation:** v35/v36 workflow generators available when relevant")
    out.append(f"**Release gate:** {gate} — {RELEASE_CRITERIA.get(gate, '')}")
    out.append(f"**Engineering score:** {base_score}/100")
    out.append(f"**Input maturity:** {maturity}/100 — {maturity_label}")
    out.append(f"**Evidence level:** {evidence_level} ({evidence_pts}/100)\n")

    if project_lines:
        out.append("#### Project frame")
        out.extend(project_lines)
        out.append("")

    out.append(knowledge_graph_markdown(gctx))
    out.append("")

    if calc_notes:
        out.append("#### Deterministic checks / validators")
        for note in calc_notes:
            out.append(f"- {note}")
        out.append("")

    out.append("#### Risk matrix")
    out.append("| Area | Risk | Engineering reason |")
    out.append("|---|---:|---|")
    for area, level, reason in risks:
        out.append(f"| {area} | {level} | {reason} |")
    out.append("")

    out.append("#### Multi-agent Engineering Board")
    out.append("| Reviewer | Stance | Review focus | Note |")
    out.append("|---|---|---|---|")
    for r in board:
        out.append(f"| {r['agent']} | {r['stance']} | {r['focus']} | {r['note']} |")
    out.append("")

    out.append("#### Engineering assessment")
    if routed == "Manufacturing / DFM":
        out.extend([
            "1. Treat this as a **pre-tooling engineering review** until CAD, material grade, nominal wall thickness, production volume and CTQ dimensions are known.",
            "2. Use the graph links to connect material, wall thickness, ribs/bosses, draft, shrinkage and tooling into one risk model, not isolated checklist items.",
            "3. High-value actions: section CAD for wall uniformity, define draft/ejection, rationalize bosses/ribs, review parting line/gates and create first-article inspection criteria.",
            "4. If volume is high, add mold-flow/tooling review and capability planning before release.",
        ])
    elif routed == "Simulation / FEA":
        out.extend([
            "1. Freeze the simulation objective before model setup: strength, stiffness, modal, buckling, fatigue or thermal behavior.",
            "2. Review loads, supports, contacts and material law before mesh refinement; wrong boundary conditions cannot be fixed by a fine mesh.",
            "3. Require convergence and validation evidence before using the result for design release.",
        ])
    elif routed == "CFD / Thermal":
        out.extend([
            "1. Calculate flow regime and boundary-condition completeness before choosing a turbulence or thermal model.",
            "2. Require mass/energy balance, residual monitoring, mesh/y+ checks and validation against analytical or test data.",
        ])
    elif routed == "CAD / SolidWorks":
        out.extend([
            "1. Treat CAD automation as a controlled engineering workflow: input validation, safe output paths, overwrite policy and error handling.",
            "2. Generate macro code only after defining document type, export formats, file naming, revision policy and test files.",
        ])
    elif routed == "Materials Selection":
        out.extend([
            "1. Select materials from functional requirements and constraints, not from a single property.",
            "2. Screen by stiffness, strength, toughness, temperature, chemical exposure, process compatibility, cost and supply availability.",
        ])
    else:
        out.extend([
            "1. Frame the problem by function, loads, material, process, interfaces, failure modes and validation evidence.",
            "2. Move from qualitative review to calculation-backed decisions as inputs mature.",
        ])
    out.append("")

    if missing:
        out.append("#### Missing inputs to close")
        for m in missing:
            out.append(f"- {m}")
        out.append("")

    out.append("#### Ranked next actions")
    actions = [
        "Complete the project profile: part function, material, process, production volume, CTQs and target environment.",
        "Attach legal datasheets, drawings, design guides or previous reports to the Reference Vault with metadata.",
        "Run the relevant calculator/validator once numeric inputs are available.",
        "Generate a report package for design/manufacturing/quality review after the evidence level reaches Level 3+.",
    ]
    if routed == "CAD / SolidWorks":
        actions.insert(2, "Use CAD Studio to generate and validate a .bas macro skeleton before running on copied SolidWorks files.")
    if routed in {"Simulation / FEA", "CFD / Thermal"}:
        actions.insert(2, "Use Simulation Studio to generate setup scoring, starter scripts and validation checklist.")
    for i, action in enumerate(actions, 1):
        out.append(f"{i}. {action}")
    out.append("")

    out.append("#### Internal citations")
    if ref_sources:
        for idx, src in enumerate(ref_sources[:6], 1):
            title = src.get("source_title") or "Uploaded reference"
            file = src.get("source_file") or "reference"
            status = src.get("approval_status") or "Draft"
            score_s = src.get("score", 0)
            out.append(f"- [R{idx}] {title} — `{file}` · status: {status} · retrieval score: {score_s}")
    if global_sources:
        for idx, src in enumerate(global_sources[:7], 1):
            out.append(f"- [K{idx}] {src.get('title')} — `knowledge_packs/{src.get('source_path')}` · retrieval score: {src.get('score')}")
    if not ref_sources and not global_sources:
        out.append("- No specific internal source was retrieved; add references or improve project metadata.")
    return "\n".join(out)

# -----------------------------------------------------------------------------
# Reports and exports
# -----------------------------------------------------------------------------

def make_project_markdown(project: Optional[Dict[str, Any]], memory: List[Dict[str, Any]], refs: List[Dict[str, Any]]) -> str:
    lines = [f"# MechAI Pro Project Report", "", f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]
    if project:
        lines.append("## Project Profile")
        for k in ["name", "project_type", "part_type", "material", "process", "manufacturing_method", "annual_volume", "target_use", "status"]:
            if project.get(k):
                lines.append(f"- **{k.replace('_', ' ').title()}:** {project.get(k)}")
        lines.append("")
    lines.append("## References")
    if refs:
        for r in refs:
            lines.append(f"- {r.get('title')} ({r.get('source_type')}, {r.get('approval_status')}, rev {r.get('revision') or '-'})")
    else:
        lines.append("- No uploaded references.")
    lines.append("")
    lines.append("## Project Memory")
    for m in memory[:80]:
        lines.append(f"### {m.get('memory_type','memory')} — {m.get('title','')}")
        lines.append(m.get("content", "")[:2500])
        lines.append("")
    return "\n".join(lines)


def bytes_docx(markdown: str) -> Optional[bytes]:
    if Document is None:
        return None
    doc = Document()
    for line in markdown.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def bytes_pdf(text: str) -> Optional[bytes]:
    if canvas is None or A4 is None:
        return None
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    w, h = A4
    y = h - 40
    c.setFont("Helvetica", 9)
    for raw in text.splitlines():
        for line in textwrap.wrap(raw, 95) or [""]:
            if y < 40:
                c.showPage()
                c.setFont("Helvetica", 9)
                y = h - 40
            c.drawString(36, y, line[:130])
            y -= 12
    c.save()
    return bio.getvalue()


def bytes_xlsx(project: Optional[Dict[str, Any]], memory: List[Dict[str, Any]], refs: List[Dict[str, Any]]) -> Optional[bytes]:
    if Workbook is None:
        return None
    wb = Workbook()
    ws = wb.active
    ws.title = "Project"
    ws.append(["Field", "Value"])
    if project:
        for k, v in project.items():
            if isinstance(v, (str, int, float)) or v is None:
                ws.append([k, v])
    wr = wb.create_sheet("References")
    wr.append(["Title", "Source Type", "Approval", "Confidentiality", "Revision", "Quality"])
    for r in refs:
        wr.append([r.get("title"), r.get("source_type"), r.get("approval_status"), r.get("confidentiality"), r.get("revision"), r.get("source_quality")])
    wm = wb.create_sheet("Memory")
    wm.append(["Type", "Title", "Content", "Created"])
    for m in memory:
        wm.append([m.get("memory_type"), m.get("title"), m.get("content"), m.get("created_at")])
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()

# -----------------------------------------------------------------------------
# Sidebar / Auth / Workspace / Project
# -----------------------------------------------------------------------------

if "messages" not in st.session_state:
    st.session_state.messages = []
if "auth_user" not in st.session_state:
    st.session_state.auth_user = None
if "current_workspace" not in st.session_state:
    st.session_state.current_workspace = None
if "current_project" not in st.session_state:
    st.session_state.current_project = None
if "selected_workspace_mode" not in st.session_state:
    st.session_state.selected_workspace_mode = "General engineering"


def default_project_payload() -> Dict[str, Any]:
    return {
        "name": "RD_Lab",
        "project_type": "General engineering",
        "part_type": "",
        "material": "",
        "process": "",
        "manufacturing_method": "",
        "annual_volume": "",
        "target_use": "",
        "status": "Active",
        "metadata": {"created_by": "auto_onboarding", "purpose": "zero_friction_start"},
    }


def ensure_default_workspace_and_project(backend: SupabaseBackend, user: Dict[str, Any]) -> Tuple[Optional[Dict[str, Any]], Optional[Dict[str, Any]], Optional[str]]:
    """Create/select a default workspace and project so a signed-in engineer can start chatting immediately."""
    try:
        workspaces = backend.list_workspaces_for_user(user)
        workspace_ids = [w.get("id") for w in workspaces]

        if st.session_state.current_workspace and st.session_state.current_workspace.get("id") in workspace_ids:
            workspace = st.session_state.current_workspace
        elif workspaces:
            workspace = workspaces[0]
            st.session_state.current_workspace = workspace
        else:
            workspace = backend.create_workspace("Personal Engineering Workspace", "Personal", user)
            if not workspace:
                return None, None, "Could not create a default workspace. Check Supabase database permissions."
            st.session_state.current_workspace = workspace

        projects = backend.list_projects(workspace["id"]) if workspace else []
        project_ids = [p.get("id") for p in projects]

        if st.session_state.current_project and st.session_state.current_project.get("id") in project_ids:
            project = st.session_state.current_project
        elif projects:
            project = projects[0]
            st.session_state.current_project = project
        else:
            project = backend.create_project(workspace["id"], user, default_project_payload())
            if not project:
                return workspace, None, "Could not create a default project. Check Supabase database permissions."
            st.session_state.current_project = project

        return workspace, project, None
    except Exception as exc:
        return None, None, f"Auto-start failed: {exc}"


status = backend.status()

with st.sidebar:
    st.markdown("### ⚙️ MechAI Pro")
    st.caption("Universal Mechanical Engineering OS")
    st.markdown(f"<span class='mech-pill'>Build {APP_VERSION}</span>", unsafe_allow_html=True)
    st.markdown("---")

    with st.expander("Storage / Auth status", expanded=False):
        st.write(f"Supabase: **{status.message}**")
        st.write(f"URL: {'✅' if status.url_present else '❌'}")
        st.write(f"Anon key: {'✅' if status.anon_present else '❌'}")
        st.write(f"Service key: {'✅' if status.service_present else '⚠️ optional but recommended'}")
        st.write(f"Bucket: `{status.bucket}`")
        st.caption("Do not upload confidential data unless your Supabase project and policies are configured correctly.")

    if not status.configured:
        st.error("Supabase is not configured. Add SUPABASE_URL and SUPABASE_ANON_KEY to Streamlit Secrets. This build needs Supabase for accounts, persistent references and project storage.")
    else:
        if st.session_state.auth_user is None:
            st.markdown("#### Account")
            auth_tab = st.radio("Mode", ["Sign in", "Create account"], horizontal=True, label_visibility="collapsed")
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            if auth_tab == "Sign in":
                if st.button("Sign in", use_container_width=True, disabled=not email or not password):
                    ok, msg, user = backend.sign_in(email, password)
                    if ok and user:
                        st.session_state.auth_user = {"id": user["id"], "email": user["email"]}
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
            else:
                if st.button("Create account", use_container_width=True, disabled=not email or not password):
                    ok, msg = backend.sign_up(email, password)
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)
        else:
            user = st.session_state.auth_user
            if st.session_state.current_workspace is None or st.session_state.current_project is None:
                ws0, pr0, auto_msg = ensure_default_workspace_and_project(backend, user)
                if auto_msg:
                    st.warning(auto_msg)
            st.markdown(f"#### Signed in")
            st.caption(user.get("email"))
            if st.button("Sign out", use_container_width=True):
                st.session_state.auth_user = None
                st.session_state.current_workspace = None
                st.session_state.current_project = None
                st.session_state.messages = []
                st.rerun()

            st.markdown("---")
            st.markdown("#### Workspace")
            workspaces = backend.list_workspaces_for_user(user)
            if not workspaces:
                st.info("Create your first workspace.")
            with st.expander("+ Create workspace", expanded=not workspaces):
                w_name = st.text_input("Workspace name", value="Personal Engineering Workspace")
                w_type = st.selectbox("Workspace type", ["Personal", "Team", "Company", "Public demo"])
                if st.button("Create workspace", use_container_width=True):
                    row = backend.create_workspace(w_name, w_type, user)
                    if row:
                        st.session_state.current_workspace = row
                        st.success("Workspace created.")
                        st.rerun()
            workspaces = backend.list_workspaces_for_user(user)
            if workspaces:
                ids = [w["id"] for w in workspaces]
                default_idx = 0
                if st.session_state.current_workspace and st.session_state.current_workspace.get("id") in ids:
                    default_idx = ids.index(st.session_state.current_workspace["id"])
                choice = st.selectbox("Select workspace", workspaces, index=default_idx, format_func=lambda w: f"{w.get('name')} · {w.get('workspace_type','')}")
                st.session_state.current_workspace = choice
                role = backend.get_role(choice["id"], user)
                st.caption(f"Role: {role}")

                if role in {"Owner", "Admin"}:
                    with st.expander("Members / permissions", expanded=False):
                        mems = backend.list_members(choice["id"])
                        for m in mems:
                            st.caption(f"{m.get('email')} — {m.get('role')}")
                        new_email = st.text_input("Invite/member email")
                        new_role = st.selectbox("Role", ROLES[1:])
                        if st.button("Add member", use_container_width=True, disabled=not new_email):
                            ok, msg = backend.add_member(choice["id"], new_email, new_role)
                            if ok:
                                st.success(msg)
                            else:
                                st.error(msg)

                st.markdown("#### Project")
                projects = backend.list_projects(choice["id"])
                with st.expander("+ Create engineering project", expanded=not projects):
                    p_name = st.text_input("Project name", value="Injection Molded Cover")
                    p_type = st.selectbox("Project type", PROJECT_TYPES)
                    part_type = st.text_input("Part type", value="Enclosure / cover")
                    material = st.text_input("Material", value="")
                    process = st.text_input("Process", value="Injection molding")
                    method = st.text_input("Manufacturing method", value="Injection molding")
                    vol = st.text_input("Annual volume", value="")
                    target = st.text_input("Target use", value="")
                    if st.button("Create project", use_container_width=True, disabled=not p_name):
                        row = backend.create_project(choice["id"], user, {
                            "name": p_name, "project_type": p_type, "part_type": part_type,
                            "material": material, "process": process, "manufacturing_method": method,
                            "annual_volume": vol, "target_use": target, "status": "Active", "metadata": {},
                        })
                        if row:
                            st.session_state.current_project = row
                            st.session_state.messages = []
                            st.success("Project created.")
                            st.rerun()
                projects = backend.list_projects(choice["id"])
                if projects:
                    pids = [p["id"] for p in projects]
                    default_p = 0
                    if st.session_state.current_project and st.session_state.current_project.get("id") in pids:
                        default_p = pids.index(st.session_state.current_project["id"])
                    proj = st.selectbox("Select project", projects, index=default_p, format_func=lambda p: f"{p.get('name')} · {p.get('project_type')}")
                    st.session_state.current_project = proj

                st.markdown("#### Engineering mode")
                st.session_state.selected_workspace_mode = st.selectbox("Workspace", list(WORKSPACES.keys()), index=list(WORKSPACES.keys()).index(st.session_state.selected_workspace_mode))

                if st.button("Clear chat", use_container_width=True):
                    st.session_state.messages = []
                    st.rerun()

# -----------------------------------------------------------------------------
# Main UI
# -----------------------------------------------------------------------------

st.markdown("### MechAI Pro — Universal Mechanical Engineering OS")
st.caption("Knowledge Graph · Multi-Agent Review Board · Calculators Pro · CAD/SolidWorks · Simulation Studio · Persistent Reference Vault")

user = st.session_state.auth_user
workspace_row = st.session_state.current_workspace
project = st.session_state.current_project
selected_mode = st.session_state.selected_workspace_mode

if not status.configured:
    st.stop()
if user is None:
    st.info("Sign in or create an account from the sidebar to use persistent projects, references, permissions and team workspaces.")
    st.stop()
if workspace_row is None or project is None:
    st.info("Preparing your default engineering workspace and project. If this message remains, click the button below.")
    if st.button("Start MechAI workspace", type="primary"):
        ws0, pr0, auto_msg = ensure_default_workspace_and_project(backend, user)
        if auto_msg:
            st.error(auto_msg)
        else:
            st.success("Workspace and project are ready.")
            st.rerun()
    st.stop()

workspace_id = workspace_row["id"]
project_id = project["id"]
role = backend.get_role(workspace_id, user)
can_upload = role in {"Owner", "Admin", "Engineer"}
can_manage = role in {"Owner", "Admin"}

# Top dashboard
cols = st.columns([1.2, 1, 1, 1])
with cols[0]:
    st.markdown(f"<div class='mech-card'><b>{html.escape(project.get('name','Project'))}</b><br><span class='mech-small'>{html.escape(project.get('project_type',''))}</span></div>", unsafe_allow_html=True)
with cols[1]:
    refs_count = len(backend.list_references(workspace_id, project_id))
    st.markdown(f"<div class='mech-card'><b>{refs_count}</b><br><span class='mech-small'>Reference sources</span></div>", unsafe_allow_html=True)
with cols[2]:
    mem_count = len(backend.list_memory(workspace_id, project_id, 300))
    st.markdown(f"<div class='mech-card'><b>{mem_count}</b><br><span class='mech-small'>Memory records</span></div>", unsafe_allow_html=True)
with cols[3]:
    st.markdown(f"<div class='mech-card'><b>{role}</b><br><span class='mech-small'>Workspace role</span></div>", unsafe_allow_html=True)

# Tabs
chat_tab, graph_tab, board_tab, vault_tab, templates_tab, calculators_tab, cad_tab, sim_tab, reports_tab, admin_tab = st.tabs([
    "Chat", "Knowledge Graph", "Review Board", "Reference Vault", "Templates", "Calculators Pro", "CAD Studio", "Simulation Studio", "Reports", "Admin / Quality",
])


with graph_tab:
    st.markdown("#### v31 Mechanical Knowledge Graph")
    st.caption("This graph links parts, materials, manufacturing processes, failure modes, validation methods and cost/quality drivers.")
    graph_question = st.text_input("Graph query", value=f"{project.get('part_type','cover')} {project.get('material','')} {project.get('process','')}")
    gctx = graph_context(graph_question, project, selected_mode)
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("Detected graph nodes", len(gctx.get("entities", [])))
    with c2:
        st.metric("Linked concepts", len(gctx.get("linked_concepts", [])))
    with c3:
        st.metric("Failure modes", len(gctx.get("failure_modes", [])))
    st.markdown(knowledge_graph_markdown(gctx))
    st.markdown("#### Graph node library")
    for node, data in GRAPH_LIBRARY.items():
        with st.expander(f"{node} · {data.get('category','concept')}"):
            st.write("Aliases:", ", ".join(data.get("aliases", [])))
            st.write("Linked concepts:", ", ".join(data.get("links", [])))
            st.write("Failure/risk modes:", ", ".join(data.get("risks", [])))
            st.write("Validation methods:", ", ".join(data.get("validation", [])))

with board_tab:
    st.markdown("#### v32–v33 Multi-Agent Engineering Review Board")
    st.caption("The board converts the current project and question into review perspectives, risk stance and release-gate logic.")
    board_prompt = st.text_area("Review board prompt", value="Create a DFM review for an injection molded plastic cover.", height=90)
    routed_board = route_workspace(board_prompt, selected_mode)
    miss_board = detect_missing_inputs(board_prompt, project, routed_board)
    risks_board = risk_matrix(board_prompt, routed_board, miss_board)
    gctx_board = graph_context(board_prompt, project, routed_board)
    bscore = score_from_risks(risks_board, miss_board)
    eval_level, eval_pts = evidence_score([], search_global_knowledge(board_prompt, routed_board, project, top_k=5), deterministic_checks_v36(board_prompt, project))
    gate = release_gate_decision(bscore, miss_board, risks_board, eval_level)
    st.markdown(f"**Routed workspace:** {routed_board}")
    st.markdown(f"**Release gate:** {gate}")
    st.markdown(f"**Engineering score:** {bscore}/100")
    st.markdown(f"**Evidence:** {eval_level} ({eval_pts}/100)")
    st.markdown("##### Board reviewers")
    for row in board_review(board_prompt, routed_board, project, miss_board, risks_board, gctx_board):
        st.markdown(f"- **{row['agent']}** — *{row['stance']}*: {row['note']}")
    st.markdown("##### Risk matrix")
    st.table([{"Area": a, "Risk": l, "Reason": r} for a, l, r in risks_board])


with vault_tab:
    st.markdown("#### Universal Reference Vault")
    st.warning("Upload only files you have the right to use. Do not upload confidential/customer-sensitive files to a public demo workspace.")
    if not can_upload:
        st.info("Your role can view references but cannot upload new references.")
    with st.form("reference_upload_form", clear_on_submit=True):
        up = st.file_uploader("Upload reference", type=["pdf", "txt", "md", "csv"], disabled=not can_upload)
        c1, c2, c3 = st.columns(3)
        with c1:
            title = st.text_input("Title")
            ref_workspace = st.selectbox("Workspace scope", list(WORKSPACES.keys()), index=list(WORKSPACES.keys()).index(selected_mode))
            source_type = st.selectbox("Source type", SOURCE_TYPES)
        with c2:
            conf = st.selectbox("Confidentiality", CONFIDENTIALITY, index=1)
            approval = st.selectbox("Approval status", APPROVAL_STATUS, index=0)
            revision = st.text_input("Revision", value="A")
        with c3:
            tags = st.text_input("Tags", value="")
            legal_note = st.text_area("Legal / usage note", value="I have the right to use this file for engineering reference inside this workspace.", height=90)
            attach_to_project = st.checkbox("Attach to current project", value=True)
        confirm = st.checkbox("I confirm this file is legal to upload and does not violate copyright, confidentiality, customer or company restrictions.")
        submitted = st.form_submit_button("Ingest reference", disabled=not can_upload or not confirm or up is None)

    if submitted and up is not None:
        data = up.getvalue()
        file_hash = sha256_bytes(data)
        suffix = Path(up.name).suffix.lower()
        if suffix not in UPLOAD_EXTENSIONS:
            st.error("Unsupported file type.")
        else:
            dup = backend.find_duplicate(workspace_id, file_hash)
            extracted = read_uploaded_text(up.name, data)
            chunks = chunk_text(extracted)
            quality = source_quality_score(source_type, approval, conf, legal_note)
            ref_id = str(uuid.uuid4())
            storage_path = f"{workspace_id}/{project_id if attach_to_project else 'workspace'}/{ref_id}_{safe_slug(up.name)}{suffix}"
            content_type = mimetypes.guess_type(up.name)[0] or "application/octet-stream"
            if dup:
                st.warning(f"Duplicate detected: this file hash already exists as `{dup.get('title')}`. Metadata will reference duplicate source.")
                storage_path = dup.get("storage_path") or storage_path
                uploaded_ok = True
            else:
                uploaded_ok, upload_msg = backend.upload_reference_file(storage_path, data, content_type)
                if not uploaded_ok:
                    st.error(upload_msg)
                    st.stop()
            ref_row = {
                "id": ref_id,
                "workspace_id": workspace_id,
                "project_id": project_id if attach_to_project else None,
                "uploaded_by": user["id"],
                "title": title or Path(up.name).stem,
                "workspace_scope": ref_workspace,
                "source_type": source_type,
                "confidentiality": conf,
                "revision": revision,
                "tags": [t.strip() for t in tags.split(",") if t.strip()],
                "legal_note": legal_note,
                "approval_status": approval,
                "source_quality": quality,
                "file_name": up.name,
                "file_ext": suffix,
                "file_size": len(data),
                "file_hash": file_hash,
                "storage_bucket": backend.bucket,
                "storage_path": storage_path,
                "duplicate_of": dup.get("id") if dup else None,
                "extracted_chars": len(extracted),
                "chunk_count": len(chunks),
                "created_at": now_iso(),
            }
            inserted = backend.insert_reference(ref_row)
            if inserted:
                chunk_rows = []
                for i, ch in enumerate(chunks):
                    chunk_rows.append({
                        "reference_id": ref_id,
                        "workspace_id": workspace_id,
                        "project_id": project_id if attach_to_project else None,
                        "chunk_index": i,
                        "chunk_text": ch,
                        "token_estimate": estimate_tokens(ch),
                        "source_quality": quality,
                        "created_at": now_iso(),
                    })
                backend.insert_chunks(chunk_rows)
                st.success(f"Reference ingested: {len(chunks)} chunks, quality score {quality}.")
                backend.add_memory(workspace_id, project_id, user, "reference_upload", title or up.name, f"Uploaded reference `{up.name}` with {len(chunks)} chunks, approval={approval}, confidentiality={conf}.", {"reference_id": ref_id})

    st.markdown("#### Search references")
    search_ref = st.text_input("Search title/tags/source")
    refs = backend.list_references(workspace_id, project_id, search_ref)
    if refs:
        for r in refs[:40]:
            dup_text = " · duplicate" if r.get("duplicate_of") else ""
            st.markdown(
                f"<div class='mech-card'><b>{html.escape(r.get('title') or '')}</b>{dup_text}<br>"
                f"<span class='mech-small'>{html.escape(r.get('source_type') or '')} · {html.escape(r.get('approval_status') or '')} · {html.escape(r.get('confidentiality') or '')} · quality {r.get('source_quality')}</span><br>"
                f"<span class='mech-small'>File: {html.escape(r.get('file_name') or '')} · chunks {r.get('chunk_count')} · rev {html.escape(r.get('revision') or '')}</span></div>",
                unsafe_allow_html=True,
            )
    else:
        st.caption("No references found yet.")

with templates_tab:
    st.markdown("#### Engineering Templates Library")
    templates = {
        "DFM Review Template": "Part/process/material/volume, wall thickness, draft, ribs/bosses, tooling, tolerances, assembly, quality plan, cost drivers.",
        "FEA Setup Review Template": "Objective, materials, loads, constraints, contacts, element type, mesh convergence, validation and failure criteria.",
        "CFD Setup Review Template": "Domain, fluid, flow regime, boundary conditions, turbulence, mesh/y+, convergence, mass/energy balance, validation.",
        "Material Selection Matrix": "Functional requirements, candidates, screening criteria, risks, manufacturing compatibility, cost and availability.",
        "Design Review Checklist": "Function, loads, interfaces, tolerances, manufacturing, validation, serviceability, safety and compliance.",
        "FMEA Template": "Function, failure mode, effects, causes, controls, severity, occurrence, detection, RPN, actions.",
        "DVP&R Template": "Requirement, test method, sample size, acceptance criteria, owner, status, evidence.",
        "Cost Reduction Template": "Material, process, assembly, inspection, scrap, cycle time, standardization, supplier alternatives.",
        "SolidWorks Macro Request Template": "Document type, input folder, export formats, naming rule, overwrite policy, error handling, output report.",
    }
    for name, desc in templates.items():
        with st.expander(name):
            st.write(desc)
            if st.button(f"Use {name}", key=f"tmpl_{name}"):
                st.session_state.messages.append({"role": "user", "content": f"Use the {name} for project {project.get('name')}."})
                st.rerun()

with calculators_tab:
    st.markdown("#### Engineering Calculators Pro")
    calc = st.selectbox("Calculator", ["Reynolds number", "Beam center-load deflection", "Shaft torsion", "Bearing L10 life", "Fastener preload", "Spring compression", "Heat transfer convection", "Pressure drop quick estimate", "Sheet metal bend allowance", "Injection molding sanity checks"])
    if calc == "Reynolds number":
        rho = st.number_input("Density ρ kg/m³", value=1000.0)
        v = st.number_input("Velocity V m/s", value=1.0)
        d = st.number_input("Hydraulic diameter D m", value=0.02, format="%.5f")
        mu = st.number_input("Dynamic viscosity μ Pa·s", value=0.001, format="%.6f")
        re_val = rho * v * d / mu if mu else 0
        st.metric("Reynolds number", f"{re_val:,.0f}")
        st.write("Regime:", "Laminar" if re_val < 2300 else "Transitional" if re_val < 4000 else "Turbulent")
    elif calc == "Beam center-load deflection":
        P = st.number_input("Load P N", value=100.0)
        L = st.number_input("Span L m", value=1.0)
        E = st.number_input("Elastic modulus E Pa", value=2.1e11, format="%.4e")
        I = st.number_input("Second moment I m⁴", value=1e-8, format="%.4e")
        sigma = st.number_input("Max moment arm c/I not included; use bending stress separately", value=0.0)
        delta = P * L ** 3 / (48 * E * I) if E and I else 0
        st.metric("Deflection δ", f"{delta:.6g} m")
    elif calc == "Shaft torsion":
        T = st.number_input("Torque T N·m", value=50.0)
        d = st.number_input("Diameter d m", value=0.02, format="%.5f")
        tau = 16 * T / (math.pi * d ** 3) if d else 0
        st.metric("Max torsional shear τ", f"{tau/1e6:.2f} MPa")
    elif calc == "Bearing L10 life":
        C = st.number_input("Dynamic rating C N", value=10000.0)
        P = st.number_input("Equivalent load P N", value=2000.0)
        n = st.number_input("Speed rpm", value=1000.0)
        pexp = st.selectbox("Bearing exponent", [3.0, 10/3], format_func=lambda x: "3 ball" if x == 3.0 else "10/3 roller")
        l10h = (1e6 / (60 * n)) * ((C / P) ** pexp) if P and n else 0
        st.metric("L10 life", f"{l10h:,.0f} h")
    elif calc == "Fastener preload":
        proof = st.number_input("Proof load or allowable clamp load N", value=10000.0)
        pct = st.number_input("Target preload fraction", value=0.75)
        st.metric("Target preload", f"{proof * pct:,.0f} N")
        st.caption("Use verified fastener grade, thread condition, lubrication and joint stiffness before release.")
    elif calc == "Spring compression":
        kspring = st.number_input("Spring rate k N/mm", value=10.0)
        travel = st.number_input("Compression travel mm", value=5.0)
        st.metric("Spring force", f"{kspring * travel:.2f} N")
    elif calc == "Heat transfer convection":
        htc = st.number_input("h W/m²K", value=25.0)
        area = st.number_input("Area m²", value=0.1)
        dt = st.number_input("ΔT K", value=20.0)
        st.metric("Q = hAΔT", f"{htc * area * dt:.2f} W")
    elif calc == "Pressure drop quick estimate":
        f = st.number_input("Darcy friction factor f", value=0.02)
        Lp = st.number_input("Pipe length L m", value=10.0)
        Dp = st.number_input("Diameter D m", value=0.02, format="%.5f")
        rho_p = st.number_input("Density ρ kg/m³", value=1000.0)
        vel_p = st.number_input("Velocity V m/s", value=1.0)
        dp = f * (Lp / Dp) * rho_p * vel_p ** 2 / 2 if Dp else 0
        st.metric("Darcy-Weisbach ΔP", f"{dp:,.1f} Pa")
    elif calc == "Sheet metal bend allowance":
        angle = st.number_input("Bend angle degrees", value=90.0)
        r = st.number_input("Inside radius mm", value=1.0)
        t = st.number_input("Thickness mm", value=1.0)
        k = st.number_input("K-factor", value=0.33)
        ba = (math.pi / 180) * angle * (r + k * t)
        st.metric("Bend allowance", f"{ba:.3f} mm")
    else:
        wall = st.number_input("Nominal wall thickness mm", value=2.5)
        rib = st.number_input("Rib thickness mm", value=1.2)
        boss = st.number_input("Boss wall/equivalent thickness mm", value=2.5)
        st.write("Rib/wall ratio:", round(rib / wall, 2) if wall else "-")
        st.write("Boss/wall ratio:", round(boss / wall, 2) if wall else "-")
        st.write("Risk note:", "Higher sink risk" if wall and boss / wall > 0.8 else "Preliminary ratio acceptable; validate with material/tooling.")

with cad_tab:
    st.markdown("#### CAD / SolidWorks Automation Studio")
    macro_req = st.text_area("Macro/workflow request", value="Export active SolidWorks document to STEP and DXF into a MechAI_Exports folder.", height=100)
    macro_code = generate_vba_macro(macro_req)
    st.code(macro_code, language="vbnet")
    st.download_button("Download .bas macro", data=macro_code.encode("utf-8"), file_name="mechai_solidworks_export.bas", mime="text/plain")
    st.info("This is a safe starter skeleton. Actual SolidWorks execution requires a local Windows machine with SolidWorks installed and a future local bridge.")

with sim_tab:
    st.markdown("#### FEA / CFD Simulation Studio")
    sim_req = st.text_area("Simulation request", value="Create an ANSYS static structural setup plan for a bracket loaded by 2 kN.", height=100)
    c1, c2 = st.columns(2)
    with c1:
        apdl = generate_apdl(sim_req)
        st.code(apdl, language="text")
        st.download_button("Download ANSYS APDL starter .mac", data=apdl.encode("utf-8"), file_name="mechai_ansys_starter.mac")
    with c2:
        jou = generate_fluent_journal(sim_req)
        st.code(jou, language="text")
        st.download_button("Download Fluent journal starter .jou", data=jou.encode("utf-8"), file_name="mechai_fluent_starter.jou")

with reports_tab:
    st.markdown("#### Engineering Report Studio")
    memory = backend.list_memory(workspace_id, project_id, 200)
    refs = backend.list_references(workspace_id, project_id)
    md_report = make_project_markdown(project, memory, refs)
    st.download_button("Download Markdown report", md_report.encode("utf-8"), file_name=f"{safe_slug(project.get('name'))}_report.md")
    docx_b = bytes_docx(md_report)
    pdf_b = bytes_pdf(md_report)
    xlsx_b = bytes_xlsx(project, memory, refs)
    if docx_b:
        st.download_button("Download Word DOCX", docx_b, file_name=f"{safe_slug(project.get('name'))}_report.docx")
    if pdf_b:
        st.download_button("Download PDF", pdf_b, file_name=f"{safe_slug(project.get('name'))}_report.pdf")
    if xlsx_b:
        st.download_button("Download Excel XLSX", xlsx_b, file_name=f"{safe_slug(project.get('name'))}_report.xlsx")
    zip_bio = io.BytesIO()
    with zipfile.ZipFile(zip_bio, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("project_report.md", md_report)
        z.writestr("project_profile.json", json.dumps(project, indent=2, ensure_ascii=False))
        z.writestr("references.json", json.dumps(refs, indent=2, ensure_ascii=False))
        z.writestr("memory.json", json.dumps(memory, indent=2, ensure_ascii=False))
    st.download_button("Download full project package ZIP", zip_bio.getvalue(), file_name=f"{safe_slug(project.get('name'))}_package.zip")

with admin_tab:
    st.markdown("#### Quality / Admin Foundation")
    st.write("This build includes v31-v36: Knowledge Graph, Multi-Agent Review Board, Release Gates, Calculators Pro, CAD/SolidWorks Agent and Simulation Studio on top of the Supabase storage/auth/reference foundation.")
    st.markdown("- Auth: Supabase email/password")
    st.markdown("- Multi-tenant foundation: workspace + members + roles")
    st.markdown("- Persistent references: Supabase Storage + Postgres metadata/chunks")
    st.markdown("- Reference intelligence: file hash duplicate detection, quality score, approval status, revision tracking")
    st.markdown("- Retrieval: global knowledge + project/workspace references with citations")
    if can_manage:
        st.success("You have admin/owner permissions for this workspace.")
    else:
        st.info("You have non-admin permissions for this workspace.")

with chat_tab:
    st.markdown("#### Engineering Conversation")
    if not st.session_state.messages:
        st.markdown(
            "<div class='mech-card'><b>Start with a project-aware engineering question.</b><br>"
            "<span class='mech-small'>Example: Create a DFM review for this injection molded ABS cover using the project references.</span></div>",
            unsafe_allow_html=True,
        )
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    prompt = st.chat_input("Ask MechAI about this engineering project...")
    if prompt:
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        with st.chat_message("assistant"):
            with st.spinner("Running internal retrieval, project memory and engineering reasoning..."):
                routed = route_workspace(prompt, selected_mode)
                global_sources = search_global_knowledge(prompt, routed, project, top_k=8)
                ref_sources = backend.search_reference_chunks(workspace_id, project_id, prompt, project, top_k=7)
                answer = build_engineering_answer(prompt, routed, project, global_sources, ref_sources)
                st.markdown(answer)
                st.session_state.messages.append({"role": "assistant", "content": answer})
                backend.add_memory(workspace_id, project_id, user, "question", prompt[:160], prompt, {"routed_workspace": routed})
                backend.add_memory(workspace_id, project_id, user, "answer", f"Answer: {prompt[:80]}", answer, {
                    "routed_workspace": routed,
                    "global_sources": [s.get("source_path") for s in global_sources[:5]],
                    "reference_sources": [s.get("reference_id") for s in ref_sources[:5]],
                })
